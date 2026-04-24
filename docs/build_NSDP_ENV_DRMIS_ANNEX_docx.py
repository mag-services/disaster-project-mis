#!/usr/bin/env python3
"""
Build docs/NSDP_ENV_DRMIS_ANNEX.docx from structured indicator rows.

Requires: pip install python-docx

Usage (from repo root):
  python3 -m venv .venv && .venv/bin/pip install python-docx
  .venv/bin/python docs/build_NSDP_ENV_DRMIS_ANNEX_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = Path(__file__).resolve().parent / "NSDP_ENV_DRMIS_ANNEX.docx"

# (code, indicator_short, drmis_role, whats_needed, how_to_accomplish_detail)
# Fifth column: detailed operational steps for P & S only; "—" otherwise.

def _blank() -> str:
    return "—"


ROWS_ENV1 = [
    ("ENV 1.1.1", "Households producing own food (own consumption)", "—", "Agriculture census / household survey (VNSO, MALFFB); NSDP portal update cycle; baseline refresh.", _blank()),
    ("ENV 1.1.2", "Up-to-date primary production data (ag, livestock, forestry, fisheries)", "—", "Sector administrative data pipelines to VNSO; annual + 6-monthly reporting per framework; data standards across MALFFB, Fisheries, Forestry, Biosecurity.", _blank()),
    ("ENV 1.1.3", "Average incidence of household food poverty (Food Poverty Line)", "—", "Household Income and Expenditure Survey / poverty analytics (VNSO); agreed methodology; regular publication.", _blank()),
    ("ENV 1.2.1", "National food guidelines (aelan kaikai) operational", "—", "Health + agriculture policy; guidelines completion; monitoring via surveys (proportion consuming island food).", _blank()),
    ("ENV 1.3.1", "Volume of targeted food imports substitutable domestically", "—", "Customs/trade statistics + DARD product list; consistent classification; annual compilation.", _blank()),
    ("ENV 1.3.2", "Domestic food price level", "—", "VNSO CPI / food price series; NSDP baseline survey linkage.", _blank()),
    ("ENV 1.3.3", "Food price inflation / food price index vs national CPI", "—", "VNSO price statistics; methodology documented; portal verification links working.", _blank()),
    ("ENV 1.4.1", "FADs per 1,000 population by province", "—", "MALFFB FAD registry; population denominators (VNSO); provincial breakdown.", _blank()),
    ("ENV 1.4.2", "People trained in food storage / preservation (annual)", "—", "Training management information system across MALFFB/NGOs; % preservation training reporting rule.", _blank()),
    ("ENV 1.4.3", "Land use capability mapping systems (100% provinces)", "W", "National land-use / agriculture mapping program; MoCCA or Lands + MALFFB; **optional:** host capability layers in DRMIS if adopted as operational GIS.", _blank()),
    ("ENV 1.5.1", "Producers with Agro-Met training", "W", "Agriculture census / training registers; **optional:** agro-climate tabular layers in DRMIS for planning.", _blank()),
    ("ENV 1.5.2", "Crop diversity index by province", "W", "Survey design (VNSO assistance per framework); agricultural census modules; provincial aggregation.", _blank()),
    ("ENV 1.5.3", "Farmers in integrated farming practices", "W", "Census/ag survey question set; reporting to portal.", _blank()),
]

ROWS_ENV2 = [
    ("ENV 2.1.1", "Green / blue-green growth in national policies & new projects", "—", "DSPPAC development projects database; Aid Coordination; definition of “green infrastructure”; annual project tagging.", _blank()),
    ("ENV 2.2.1", "Development projects complying with EIA requirements", "—", "DEPC EIA register; numerator/denominator agreed; portal sync.", _blank()),
    ("ENV 2.2.2", "Environmental guidelines in place (under Acts)", "—", "DEPC policy roll-out; count vs target (6 by 2030).", _blank()),
    ("ENV 2.3.1", "Imported lighting & refrigeration meeting efficiency standards", "—", "Customs / standards compliance data; market surveillance (Energy/Customs).", _blank()),
    ("ENV 2.3.2", "Households using renewable energy as main lighting source", "—", "Household survey (VNSO) or energy census; periodic update.", _blank()),
    ("ENV 2.4.1", "Controlled waste disposal facilities (by province)", "—", "Provincial infrastructure inventory (municipalities/DEPC); operational status.", _blank()),
    ("ENV 2.4.2", "Vehicles complying with emission standards", "—", "Transport / municipal roadworthy data linked to emission rules.", _blank()),
    ("ENV 2.4.3", "Implementation of National Waste Management Strategy & Action Plan", "—", "DEPC monitoring framework; % activities completed.", _blank()),
    ("ENV 2.5.1", "Targets met in National Environment Policy & Implementation Plan 2016–2030", "—", "DEPC corporate planning; annual self-assessment vs plan.", _blank()),
    ("ENV 2.5.2", "Environmental obligations in treaties met (corporate plans & reports)", "—", "Cross-government treaty implementation tracking; legal/policy focal points.", _blank()),
    ("ENV 2.6.1", "DEPC funding approved / gov environmental expenditure", "—", "MoF budget coding; NPP alignment; baseline 2016 VT amount tracked.", _blank()),
    ("ENV 2.6.2", "Environmental Trust Fund operational & deposits", "—", "Trust fund legal setup; treasury reporting.", _blank()),
]

ROWS_ENV3 = [
    (
        "ENV 3.1.1",
        "Ministries with CC & DRM policies, budgets, legislation",
        "S",
        "Whole-of-government survey; legal/policy inventory; MoF budget tagging; **DRMIS:** does not replace policy audit—supports evidence-based DRM dialogue.",
        "(1) NAB/OPM runs an annual structured questionnaire to each ministry: existence of CC/DRM policy, approved budget line, and cited legislation. (2) State Law Office or legal units maintain a master index of Acts and amendments. (3) MoF supplies evidence of budget tagging against DRM/CC codes. (4) Collate results in a single spreadsheet for NSDP portal entry (official %). (5) DRMIS (optional): upload read-only metadata or links to public documents; publish a non-authoritative tabular summary (ministry × yes/no) for workshops only.",
    ),
    (
        "ENV 3.1.2",
        "DoCC / NAB secretariat staff financed by government",
        "S",
        "Payroll / establishment data (PMO, MoF); **DRMIS:** unrelated except as national DRM IT asset.",
        "(1) Identify establishment positions for DoCC and NAB secretariat in official HR/payroll systems. (2) MoF confirms salaries charged to government vote vs donor. (3) Report numerator/denominator per NSDP methodology. DRMIS: no payroll data; optionally note hosting of DRM systems as project context in documentation only.",
    ),
    (
        "ENV 3.1.3",
        "Sector stakeholders aligned with CC & DRM policies",
        "S",
        "Formal arrangement registers; awareness surveys; **DRMIS:** operational risk/damage catalogue for DRM users only—not NSDP reporting.",
        "(1) NAB maintains a register of MOUs, TORs, and sector coordination meetings with dates and outcomes. (2) Run periodic perception or awareness surveys using agreed instruments. (3) Enter official results on the NSDP platform. (4) DRMIS: keep published hazard/damage/needs datasets current so stakeholders use one operational catalogue in meetings—does not replace survey statistics.",
    ),
    (
        "ENV 3.2.1",
        "Multi-hazard warning systems with maintenance plans (provinces)",
        "S",
        "VMGD/NDMO asset register; maintenance plans filed; **DRMIS:** situational view only—not warning system of record.",
        "(1) VMGD/NDMO maintain an asset register (sirens, radio, equipment) with location, serial, and responsible officer per province. (2) Each asset has a linked maintenance plan document ID and next service date. (3) Count provinces meeting “system + plan” rule for the portal. (4) DRMIS: optional map layer of assets if coordinates are cleared for release; read-only; operational awareness only.",
    ),
    (
        "ENV 3.2.2",
        "Population with access to technologies conveying early warnings",
        "S",
        "Survey on access (phones, radio, etc.); telecom/media partnership; **DRMIS:** optional public dashboards.",
        "(1) Define “access” (e.g. mobile ownership, radio reach) with VNSO/telecom. (2) Add module to household or dedicated ICT survey; use official population denominators. (3) Report disaggregated estimates to the NSDP platform. (4) DRMIS: optional choropleth by province if ministry publishes approved tabular extracts and MOU allows display—dashboard is illustrative, not the survey of record.",
    ),
    (
        "ENV 3.2.3",
        "Knowledge / publications in atmospheric & earth sciences",
        "S",
        "Bibliometric count; VMGD, USP, research partners; bibliographic database.",
        "(1) Agree counting rules (peer-reviewed, grey literature, national authors). (2) VMGD/USP assign focal points to tag publications annually. (3) Maintain a simple bibliographic database or spreadsheet for verification. (4) DRMIS: optional link list or tabular index of titles/URLs if ministries want public visibility—not the official count.",
    ),
    (
        "ENV 3.3.1",
        "Community support plans (coordination, PDRR)",
        "P",
        "NDMO/provincial planning registers; % communities with approved plans; **DRMIS:** host or link to plan metadata if digitised.",
        "(1) NDMO and provincial offices adopt a single register schema: community ID, island, plan status (draft/approved), approval date, event link. (2) Update after each event or annual cycle. (3) Official numerator/denominator for NSDP entered on the national platform. (4) DRMIS: create TabularDataset or vector boundaries with fields plan_id, status, date; use draft/published workflow; API for partners; RAP/community imports only as approved.",
    ),
    (
        "ENV 3.3.2",
        "Disaster/CC-affected communities with durable solutions",
        "P",
        "Post-event tracking; durable solution criteria (HLP, shelter, livelihoods); **DRMIS:** operational damage/needs layers; **official indicator values: NSDP platform.**",
        "(1) With shelter/HLP cluster, define “durable solution” criteria (e.g. permanent housing, land tenure, livelihood restored). (2) Run post-event survey rounds by community with unique IDs aligned to Area Council boundaries. (3) Track transitions from temporary to durable status over time. (4) Enter verified percentages on the NSDP platform. (5) DRMIS: ingest tabular datasets (damage, needs, financial types) and RAP batches by event; province/AC filters; publication_status and audit; GIS join optional—operational evidence only.",
    ),
    (
        "ENV 3.3.3",
        "National multi-hazard & risk maps for PDNA",
        "P",
        "**Designate official map products;** MoCCA/NDMO workflow; optional publish layers in DRMIS for operations; metadata & version control; **NSDP M&E and verification: NSDP platform only.**",
        "(1) MoCCA and NDMO designate one official “national multi-hazard / risk” map product (or suite) for PDNA use: owner, version, CRS, legend, and update cycle in writing. (2) Prepare layers (vector/raster/PMTiles) to agreed specs; clear sensitive areas. (3) Publish to DRMIS with full metadata, titiler params if raster, and publication_status=published when cleared. (4) If hosted externally, store WMS/WMTS URL + licence in DRMIS catalogue. (5) NSDP portal progress and verification remain separate from DRMIS hosting.",
    ),
    ("ENV 3.4.1", "Public schools using CC & DRR curriculum modules", "W", "MoET EMIS; school inspection data; curriculum rollout tracking.", _blank()),
    ("ENV 3.4.2", "Islands covered by CC adaptation / resilience awareness", "W", "Programme logs (NDMO, NGOs); island-level coverage denominator.", _blank()),
    ("ENV 3.5.1", "CCA/DRM spend via donor budget support through government systems", "W", "NAB project finance tracking; MoF on-budget reporting.", _blank()),
    ("ENV 3.5.2", "Accreditation to climate & disaster finance funds (e.g. AF, GCF)", "—", "NAB/DMFU accreditation milestones; external fund requirements.", _blank()),
    ("ENV 3.5.3", "Volume of external climate/disaster finance for community programmes", "—", "Project disbursement tracking; FX and programme classification.", _blank()),
]

ROWS_ENV4 = [
    (
        "ENV 4.1.1",
        "Declared physical planning areas with an approved physical plan",
        "S",
        "Lands / municipal planning registers; **DRMIS:** optional spatial overlay of declared areas vs plans.",
        "(1) Lands and municipal governments maintain authoritative registers of declared planning areas and approved physical plans with unique IDs. (2) Capture or digitise plan boundaries as GIS polygons aligned to national CRS. (3) Gap analysis: declared without plan vs plan in place. (4) Official statistics to NSDP platform. (5) DRMIS: optional VectorDataset layers (province/AC attributes); draft vs published; no replacement for legal land registry.",
    ),
    (
        "ENV 4.1.2",
        "Directives of national land use planning policy implemented",
        "S",
        "Policy implementation checklist by agency; periodic assessment.",
        "(1) DEPC/Lands agree a checklist of NUPP directives with lead agency per directive. (2) Annual self-assessment: status (not started / in progress / complete) and evidence URL or file reference. (3) Aggregate % implemented for portal reporting. (4) DRMIS: optional attachment of spatial layers cited as evidence where ministries approve publication.",
    ),
    (
        "ENV 4.2.1",
        "Declared Water Protection Zones",
        "S",
        "DoW / DEPC register; legal gazettal; **DRMIS:** optional boundary layers.",
        "(1) DoW maintains the legal list of Water Protection Zones with gazettal references and dates. (2) Digitise boundaries from legal descriptions into GIS where possible. (3) Target count (e.g. 6 zones) tracked on NSDP platform. (4) DRMIS: VectorDataset with attributes gazettal_date, status; security review before publish.",
    ),
    (
        "ENV 4.2.2",
        "GIS mapping for forest management (significant forest areas)",
        "S",
        "Forestry / Planning & Mapping Unit; integration with land capability; **DRMIS:** host forest significance layers if national GIS strategy assigns.",
        "(1) Forestry and Planning & Mapping Unit produce a national mosaic of “significant forest” consistent with national forest policy. (2) Align with land use capability classes where the framework requires. (3) Update on agreed cycle with change log. (4) DRMIS: ingest as RasterDataset (Climate) with legend, year, and is_land_cover rules; TiTiler/rescale as needed; assign ownership in metadata.",
    ),
    ("ENV 4.3.1", "Mineral extraction complying with EIA", "—", "DEPC + mining regulator compliance statistics.", _blank()),
    ("ENV 4.4.1", "National fishing fleet compliance with RFMOs", "—", "Fisheries business plan metrics; licensing and observer data.", _blank()),
    (
        "ENV 4.4.2",
        "Integrated coastal management plans developed & operational",
        "S",
        "**DRMIS:** coastal datasets and PMTiles support spatial planning; Fisheries + provincial coordination.",
        "(1) Fisheries leads the 40-plan (or current target) programme with provincial partners; each plan has adoption date and legal status. (2) Operationalise plans with assigned budgets and focal points. (3) DRMIS: ingest coastal PMTiles/vector for ICM zones where boundaries are cleared; cyclone_name and metadata for events; coordinate with community profile coastal modules if policy allows.",
    ),
    ("ENV 4.5.1", "Approved Fisheries Management Plans for designated fisheries", "—", "Fisheries department plan register.", _blank()),
    ("ENV 4.5.2", "Commercial coastal sand mining (ban in 6 sites by 2025)", "—", "MoIA / foreshore permits; enforcement reporting.", _blank()),
    (
        "ENV 4.5.3",
        "Places with detailed geo-scientific data collected",
        "S",
        "Geological survey / coastal studies inventory; **DRMIS:** optional publication of non-sensitive spatial products.",
        "(1) Build a national inventory of geo-scientific studies (points or polygons) with study name, year, lead agency, and abstract. (2) Define “100%” target per NSDP methodology. (3) DRMIS: publish only non-sensitive layers; classify sensitive studies as draft or omit; use publication_status and audit fields.",
    ),
    ("ENV 4.6.1", "Annual area of reforestation / rehabilitation", "W", "Planning & Mapping Unit database; field verification; annual reports.", _blank()),
    ("ENV 4.6.2", "Forestry licensees within quota / monitoring coverage", "W", "Forestry compliance system; REDD+ readiness linkage per national plan.", _blank()),
    ("ENV 4.7.1", "DEPC outreach & provincial offices", "W", "DEPC annual report activity counts; establishment of 6 provincial offices by 2030.", _blank()),
]

ROWS_ENV5 = [
    ("ENV 5.1.1", "NISSAP / NBSAP targets implementation", "—", "DEPC implementation tracking; % activities complete.", _blank()),
    ("ENV 5.1.2", "Environmental guidelines in place", "—", "DEPC guideline register vs target count.", _blank()),
    ("ENV 5.2.1", "Registered Community Conservation Areas (CCAs)", "—", "DEPC / conservation registry.", _blank()),
    ("ENV 5.2.2", "CCA committees implementing management plans", "—", "Site-level reporting; plan adoption tracking.", _blank()),
    ("ENV 5.3.1", "Population knowledge of local flora & fauna (traditional knowledge baseline)", "—", "Survey instrument; cultural knowledge protocols; baseline by 2020 (framework).", _blank()),
    ("ENV 5.3.2", "Threatened species legally protected (IUCN alignment)", "—", "Legislation updates; legal protection list vs IUCN Red List.", _blank()),
    ("ENV 5.4.1", "Planes & sea vessels cleared through Biosecurity", "—", "Biosecurity operational stats; total arrivals denominator.", _blank()),
    ("ENV 5.5.1", "Model schools implementing environmental programmes", "—", "MoET programme designation; per-province model school count.", _blank()),
    (
        "ENV 5.6.1",
        "Central information sharing system for environment data",
        "S",
        "**National architecture decision** (UNDP/GEF and partners); APIs; metadata catalogue; **DRMIS** as hazard/disaster node or federated publisher.",
        "(1) Government adopts a national architecture for environment data (catalogue service, metadata standard, access policy)—often led by DEPC with UNDP/GEF support. (2) Register authoritative datasets with persistent IDs and licences. (3) Expose APIs or CSW for discovery; avoid duplicate master databases. (4) DRMIS: operate as a **node** publishing hazard/disaster/climate-context datasets with API keys; participate in federated search rather than merging all environment data into one stack. (5) NDMO/MoCCA data governance for DRMIS scope.",
    ),
]

# RAP/vanuatu — sibling repo; PDNA-style tabulations (may inform ministries; not NSDP via DRMIS)
ROWS_RAP = [
    (
        "data/1c_input_baseline.csv",
        "Pre-disaster indicators by sector, Area Council, year (feeds modelled outputs).",
        "ENV 1 food/ag; ENV 3.3.x exposure — ministries own official stats.",
    ),
    (
        "data/Ex_hazard_areas.csv",
        "Cyclone intensity 2–5 per council (from VMGD track).",
        "ENV 3.2.x / 3.3.3 hazard context.",
    ),
    (
        "data/Ex_hazard_areas_MIS_import.csv",
        "Hazard export for MIS tabular import (scripts/export_hazard_for_mis.R).",
        "Operational DRMIS load — not an NSDP indicator return.",
    ),
    (
        "output/*_01*_baseline*.csv",
        "Aggregated baselines: Education, FoodSecurity, Health, Shelter, WASH, Logistics, Energy, Telecom, Gender+Protection, Business.",
        "ENV 1.4/1.5 ag; ENV 3.4 schools; sector context for dialogue.",
    ),
    (
        "output/*_02* damage, *_03* resources, *_04* financial",
        "Modelled damage, immediate response resources, unit-cost financial damage.",
        "ENV 3.3.x PDNA-style evidence class only; not portal figures.",
    ),
    (
        "output/QC_baseline_coverage.csv",
        "Row counts per sector × Area Council (QA).",
        "Coverage tracking for modellers.",
    ),
    (
        "output/area_councils_intensity.geojson",
        "AC polygons with intensity for GIS / DRMIS map layers.",
        "ENV 3.3.3 operational map input if published to DRMIS.",
    ),
]


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_rap_section(doc: Document) -> None:
    add_heading(doc, "RAP Vanuatu tabulation outputs (not NSDP reporting via DRMIS)", level=2)
    doc.add_paragraph(
        "Sibling repository RAP/vanuatu (Quarto) produces cyclone baseline, damage, resource, "
        "and financial tabulations from council-level baselines and hazard intensity. "
        "Use toward Environment pillar indicators only through line ministries and the NSDP platform."
    )
    table = doc.add_table(rows=1 + len(ROWS_RAP), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "Path / pattern"
    hdr[1].text = "Description"
    hdr[2].text = "Thematic ENV link (illustrative)"
    for cell in hdr:
        p = cell.paragraphs[0]
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
    for i, (path, desc, link) in enumerate(ROWS_RAP, start=1):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = desc
        row[2].text = link
        for c in row:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def _set_cell_font(cell, size_pt: float, bold: bool = False) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size_pt)
            if bold:
                r.bold = True


def add_table_section(doc: Document, title: str, rows: list[tuple[str, str, str, str, str]]) -> None:
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "Code"
    hdr[1].text = "Indicator (summary)"
    hdr[2].text = "DRMIS role"
    hdr[3].text = "What is needed (summary)"
    hdr[4].text = "How to accomplish — detail for P & S"
    for j, cell in enumerate(hdr):
        _set_cell_font(cell, 9 if j < 4 else 8, bold=True)
    for i, (code, ind, drmis, needed, how) in enumerate(rows, start=1):
        row = table.rows[i].cells
        row[0].text = code
        row[1].text = ind
        row[2].text = drmis
        row[3].text = needed
        row[4].text = how
        _set_cell_font(row[0], 8)
        _set_cell_font(row[1], 8)
        _set_cell_font(row[2], 8)
        _set_cell_font(row[3], 8)
        _set_cell_font(row[4], 7)
        for c in row:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def set_landscape(section) -> None:
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    set_landscape(section)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    t = doc.add_paragraph()
    t.add_run("Annex: NSDP Environment pillar — thematic crosswalk (DRMIS not M&E)").bold = True
    t.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "Purpose: Optional planning map between Environment pillar indicator codes and DRMIS capabilities. "
        "Official NSDP monitoring and evaluation—indicators, baseline/target/current, verification—is on the "
        "national NSDP platform only. DRMIS does not inform, feed, or substitute for that platform."
    )
    doc.add_paragraph(
        "Legend — DRMIS role: P = primary/strong; S = secondary/contributory; "
        "W = weak/conditional; — = out of scope unless datasets are explicitly added. "
        "The fifth column gives step-by-step operational detail for every P and S row; "
        "other rows show an em dash in that column."
    )
    doc.add_paragraph(
        "Sources: (1) NSDP M&E Framework (July 2017), DSPPAC — Environment pillar tables. "
        "(2) Environment Pillar: Full NSDP Indicator Progress Factsheet (Environment Pillar All objectives.pdf) "
        "— 63 indicators, portal status, critical observations on data gaps."
    )
    doc.add_paragraph(
        "Indicator IDs follow the factsheet-style numbering (e.g. ENV 3.3.3). "
        "Resolve any goal-title mismatches between PDFs using the live NSDP tracking portal."
    )

    add_rap_section(doc)

    add_table_section(doc, "Environment 1 — Food security & sustainable production", ROWS_ENV1)
    add_table_section(doc, "Environment 2 — Natural resource / green growth & compliance (2017 + factsheet themes)", ROWS_ENV2)
    add_table_section(doc, "Environment 3 — Climate change & disaster risk reduction", ROWS_ENV3)
    add_table_section(doc, "Environment 4 — Environmental protection, land, water, coastal, forests", ROWS_ENV4)
    add_table_section(doc, "Environment 5 — Biodiversity, biosecurity, education, data sharing", ROWS_ENV5)

    doc.add_paragraph()
    add_heading(doc, "DRMIS capabilities (operational reference only)", level=2)
    hook_table = doc.add_table(rows=6, cols=2)
    hook_table.style = "Table Grid"
    hooks = [
        ("Clusters + disaster overlay tags", "Operational hazard/exposure groupings"),
        ("Province & area council data", "Subnational views inside DRMIS"),
        ("Dataset types: baseline, damage, needs, financial", "PDNA-style ops data; not NSDP returns"),
        ("Raster / PMTiles / vector", "Spatial products; policy parallel to ENV 3/4 only"),
        ("API + integration keys", "Partner access to published DRMIS data"),
        ("Publication + audit", "DRMIS-internal traceability — not NSDP verification"),
    ]
    for i, (a, b) in enumerate(hooks):
        hook_table.rows[i].cells[0].text = a
        hook_table.rows[i].cells[1].text = b

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(
        "Generated for DRMIS repository. Regenerate with: python docs/build_NSDP_ENV_DRMIS_ANNEX_docx.py "
        "(after pip install python-docx)."
    ).italic = True
    foot.runs[0].font.size = Pt(8)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
