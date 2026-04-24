#!/usr/bin/env python3
"""
Build docs/NDMO_MOCCA_DRMIS_Briefing.docx for NDMO → MoCCA director briefing.

Requires: pip install python-docx

Usage:
  .venv/bin/python docs/build_NDMO_MOCCA_briefing_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

OUT_PATH = Path(__file__).resolve().parent / "NDMO_MOCCA_DRMIS_Briefing.docx"

# (#, product, type, why, who, what_needed, baseline_H, target_I, current_J)
# H/I/J from Environment Pillar All objectives.pdf (NSDP factsheet).
DATA_ROWS: list[tuple[str, ...]] = [
    (
        "1",
        "Gridded rainfall (monthly/seasonal/annual climatology; anomalies)",
        "Raster",
        "DRR; ENV 3. H/I/J → ENV 3.2.3 (atmospheric/earth sciences).",
        "MoCCA / Dept Climate Change; VMGD if held there.",
        "COG/GeoTIFF; CRS; metadata; cadence; licence; filename_id; focal officer.",
        "TBD",
        "Increase in publications or research proposals in atmospheric and earth sciences",
        "-",
    ),
    (
        "2",
        "Temperature (mean/min/max surfaces, anomalies)",
        "Raster",
        "Heat/health; ENV 3. H/I/J → ENV 3.2.3 (same proxy as row 1).",
        "MoCCA / VMGD",
        "As row 1; observation vs reanalysis documented.",
        "TBD",
        "Increase in publications or research proposals in atmospheric and earth sciences",
        "-",
    ),
    (
        "3",
        "Standardised drought index (SPI, SPEI)",
        "Raster or tabular",
        "Early action; ENV 1.5/3. H/I/J → ENV 1.5.1 (Agro-Met training).",
        "MoCCA (+ MALFFB / DoW validation)",
        "Methodology; monthly refresh; province aggregates if needed.",
        "TBD",
        "Ag Census infor",
        "-",
    ),
    (
        "4",
        "Seasonal climate outlook (maps or province categories)",
        "Raster or tabular",
        "Preparedness; H/I/J → ENV 3.2.1 (multi-hazard warning systems).",
        "VMGD / MoCCA",
        "Official layers only; version date; link to advisory.",
        "TBD",
        "100% of provinces with multi-hazard warning systems",
        "-",
    ),
    (
        "5",
        "Climate projections (rainfall/temp scenarios)",
        "Raster (multi-layer)",
        "Adaptation; H/I/J → ENV 3.4.2 (awareness programs).",
        "MoCCA (+ partners)",
        "Scenario labels; IPCC/study ref; public vs restricted.",
        "TBD",
        "60% of islands covered by awareness programs targeting climate change adaptation and resilience",
        "-",
    ),
    (
        "6",
        "Tropical cyclone / severe weather footprints",
        "Vector / PMTiles",
        "PDNA; H/I/J → ENV 3.2.2 (access to early-warning technologies).",
        "VMGD / NDMO (events)",
        "Event ID, time, authoritative geometry.",
        "TBD",
        "0.8",
        "0.6",
    ),
    (
        "7",
        "Flood / tsunami / landslide national reference layers",
        "Vector / Raster",
        "ENV 3.3.3 multi-hazard maps. H/I/J → ENV 3.3.3.",
        "MoCCA, NDMO, Public Works / GeoHazards",
        "Ownership; post-event updates.",
        "TBD",
        "1 national multi-hazard and risk map",
        "50",
    ),
    (
        "8",
        "Sea-level rise / coastal inundation scenarios",
        "Raster",
        "ENV 4 coastal. H/I/J → ENV 4.5.3 (proxy; no dedicated SLR indicator).",
        "MoCCA / coastal unit",
        "Datum; scenario years; Fisheries ICM.",
        "TBD",
        "100%",
        "-",
    ),
    (
        "9",
        "Land cover / land use change (national mosaic)",
        "Raster",
        "ENV 4.2.2. H/I/J → ENV 4.2.2.",
        "MoCCA / Forestry / Lands",
        "Single authoritative mosaic/year; legend.",
        "Creating precise GIS forest maps integrated with Land Use Capability data.",
        "By 2030 100% of forest areas of significance mapped",
        "-",
    ),
    (
        "10",
        "Post-event damage (verified counts by area)",
        "Tabular / vector",
        "PDNA; H/I/J → ENV 3.3.2 (durable solutions).",
        "NDMO (+ sectors)",
        "Survey rules; draft vs published.",
        "TBD",
        "60% of climate change and disaster affected communities with durable solutions",
        "-",
    ),
    (
        "11",
        "Evacuation centres / safe sites",
        "Vector",
        "Operations; H/I/J → ENV 3.3.1 (support plans).",
        "NDMO / Provincial",
        "Verified coords; seasonal refresh.",
        "TBD",
        "80% of communities have access to support plans",
        "-",
    ),
    (
        "12",
        "Water stress / hydrological drought",
        "Tabular / raster",
        "DoW + NDMO. H/I/J → ENV 4.2.1 (Water Protection Zones).",
        "DoW / MoCCA",
        "Align with rainfall/SPI.",
        "TBD",
        "6 Water Protection Zones declared",
        "-",
    ),
    (
        "13",
        "Food security / agricultural stress",
        "Tabular",
        "ENV 1. H/I/J → ENV 1.1.3 (household food poverty).",
        "MALFFB (+ VNSO official)",
        "Attributes; province map.",
        "5.7",
        "5.6",
        "-",
    ),
    (
        "14",
        "Integrated coastal management boundaries",
        "Vector",
        "ENV 4.4.2. H/I/J → ENV 4.4.2.",
        "Fisheries / MoCCA",
        "Legal status; MOU updates.",
        "TBD",
        "40 integrated coastal management plans developed and operational",
        "-",
    ),
    (
        "15",
        "Physical planning area boundaries",
        "Vector",
        "ENV 4.1.1. H/I/J → ENV 4.1.1.",
        "Lands / municipalities",
        "Official GIS export; sync schedule.",
        "TBD",
        "100% of all physical planning areas declared in or before 2016 have a physical plan in place",
        "-",
    ),
]


def set_landscape(section) -> None:
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height


def main() -> None:
    doc = Document()
    # Landscape for wide table
    sec = doc.sections[0]
    set_landscape(sec)
    sec.left_margin = Inches(0.45)
    sec.right_margin = Inches(0.45)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)

    t = doc.add_paragraph()
    r = t.add_run("DRMIS data partnership — NDMO → Ministry of Climate Change")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "Audience: Director NDMO presenting to Director General, MoCCA. "
        "Purpose: align DRMIS with MoCCA priorities and ENV 3–4 policy themes; close data gaps (e.g. rainfall rasters). "
        "M&E boundary: official NSDP indicators and baseline/target/current are on the national NSDP platform only; "
        "DRMIS does not inform or substitute for that platform. "
        "H, I, J in the table are copied from the Environment Pillar factsheet for contextual read-across only."
    )

    doc.add_heading("1. Executive summary", level=1)
    for line in [
        "DRMIS: operational published hazard, exposure, damage, climate-context map/tabular layers; province/AC; APIs—not NSDP M&E.",
        "NSDP Environment themes explain why shared national layers matter for DRM; authoritative M&E figures stay on the NSDP platform.",
        "MoCCA: natural provider for gridded climate rasters — largely missing from DRMIS today.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2. DRMIS modalities", level=1)
    modalities = [
        ("Raster", "GeoTIFF/VRT/COG; TiTiler or precomputed tiles; Climate mode."),
        ("Vector", "GeoJSON; hazards, exposure, admin boundaries."),
        ("Tabular", "Province/AC stats; damage, needs, financial types."),
        ("PMTiles", "Basemap-style hazard layers from URL."),
    ]
    m = doc.add_table(rows=1 + len(modalities), cols=2)
    m.style = "Table Grid"
    m.rows[0].cells[0].text = "Modality"
    m.rows[0].cells[1].text = "Role"
    for ri, (a, b) in enumerate(modalities, start=1):
        m.rows[ri].cells[0].text = a
        m.rows[ri].cells[1].text = b

    doc.add_heading("3. Priority data gaps (with Baseline H / Target I / Current J)", level=1)

    cols = [
        "#",
        "Data product",
        "Type",
        "Why / NSDP",
        "Who provides",
        "What is needed",
        "Baseline (H)",
        "Target (I)",
        "Current (J)",
    ]
    table = doc.add_table(rows=1 + len(DATA_ROWS), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for j, h in enumerate(cols):
        cell = table.rows[0].cells[j]
        cell.text = h
        p = cell.paragraphs[0]
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(7)

    for i, row_data in enumerate(DATA_ROWS, start=1):
        for j, val in enumerate(row_data):
            c = table.rows[i].cells[j]
            c.text = val
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7)

    doc.add_heading("4. Governance (summary)", level=1)
    g = doc.add_table(rows=6, cols=2)
    g.style = "Table Grid"
    pairs = [
        ("Data-sharing", "MOU: official layers, who publishes, embargo."),
        ("Focal points", "Named MoCCA + NDMO data officers; quarterly list."),
        ("Publication standard", "CRS, title, provider, date, licence in metadata (internal tags only—not NSDP reporting)."),
        ("Security", "Draft / published / archived in DRMIS."),
        ("M&E / NSDP", "National NSDP platform only for indicators and verification; DRMIS not M&E."),
    ]
    g.rows[0].cells[0].text = "Topic"
    g.rows[0].cells[1].text = "Action"
    for ri, (a, b) in enumerate(pairs, start=1):
        g.rows[ri].cells[0].text = a
        g.rows[ri].cells[1].text = b

    doc.add_heading("5. Ask to MoCCA (DG)", level=1)
    for line in [
        "Designate provider of record for national climate rasters (or VMGD split in writing).",
        "Approve pilot: one rainfall or SPI stack in DRMIS Climate admin with metadata.",
        "Assign focal officer + update calendar.",
        "Support operational multi-hazard mapping where MoCCA underpins layers (M&E remains on NSDP platform).",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    foot = doc.add_paragraph()
    foot.add_run(
        "Regenerate: python docs/build_NDMO_MOCCA_briefing_docx.py (pip install python-docx)."
    ).italic = True
    foot.runs[0].font.size = Pt(8)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
